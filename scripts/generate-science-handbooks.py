from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATA = ROOT / "data" / "readers.json"
DOWNLOADS = ROOT / "public" / "downloads"

COURSES = {
    "cosmos": (
        "Cosmos_and_Astronomy_Farsi_Student_Reader.docx",
        "Cosmos_and_Astronomy_Teacher_Reader.docx",
    ),
    "life": (
        "Life_Science_and_the_Body_Farsi_Student_Reader.docx",
        "Life_Science_and_the_Body_Teacher_Reader.docx",
    ),
    "earth-climate": (
        "Earth_and_Climate_Farsi_Student_Reader.docx",
        "Earth_and_Climate_Teacher_Reader.docx",
    ),
    "physical-chemical": (
        "Physics_and_Chemistry_Farsi_Student_Reader.docx",
        "Physics_and_Chemistry_Teacher_Reader.docx",
    ),
    "computing-internet": (
        "Computing_and_the_Internet_Farsi_Student_Reader.docx",
        "Computing_and_the_Internet_Teacher_Reader.docx",
    ),
    "engineering-tech": (
        "Engineering_and_Technology_Farsi_Student_Reader.docx",
        "Engineering_and_Technology_Teacher_Reader.docx",
    ),
}

# compact_reference_guide preset, with two named overrides:
# - Persian reading text uses Noto Naskh Arabic at 13.5 pt.
# - The site palette replaces the preset blue hierarchy.
INK = "172421"
GREEN = "123F38"
ORANGE = "B95429"
MUTED = "65716D"
LINE = "D8E2DC"
SOFT = "E8EFEA"
PERSIAN_FONT = "Noto Naskh Arabic"
LATIN_FONT = "Calibri"


def set_paragraph_rtl(paragraph, rtl: bool = True) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if rtl and bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    elif not rtl and bidi is not None:
        p_pr.remove(bidi)


def format_run(run, *, persian: bool, size: float, bold: bool = False, color: str = INK) -> None:
    font_name = PERSIAN_FONT if persian else LATIN_FONT
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), font_name)
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "fa-IR" if persian else "en-US")
    lang.set(qn("w:bidi"), "fa-IR" if persian else "en-US")
    if persian:
        rtl = r_pr.find(qn("w:rtl"))
        if rtl is None:
            r_pr.append(OxmlElement("w:rtl"))


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_keep_with_next(paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("صفحه ")
    format_run(run, persian=True, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, value, end])


def add_bottom_border(paragraph, color: str = LINE) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def configure_document(doc: Document, title: str, teacher: bool) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)

    doc.core_properties.title = title
    doc.core_properties.subject = "Complete Persian study reader" + (" with English teaching support" if teacher else "")
    doc.core_properties.author = "Farsi Reader Library"
    doc.core_properties.keywords = "Persian, Farsi, reader, teacher, student"

    normal = doc.styles["Normal"]
    normal.font.name = PERSIAN_FONT
    normal.font.size = Pt(13.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for style_name, size, color, before, after in (
        ("Heading 1", 16, GREEN, 18, 10),
        ("Heading 2", 13, GREEN, 14, 7),
        ("Heading 3", 12, GREEN, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = PERSIAN_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "English Gloss" not in doc.styles:
        gloss = doc.styles.add_style("English Gloss", WD_STYLE_TYPE.PARAGRAPH)
    else:
        gloss = doc.styles["English Gloss"]
    gloss.font.name = LATIN_FONT
    gloss.font.size = Pt(9.5)
    gloss.font.color.rgb = RGBColor.from_string(MUTED)
    gloss.paragraph_format.left_indent = Inches(0.24)
    gloss.paragraph_format.space_after = Pt(7)
    gloss.paragraph_format.line_spacing = 1.08
    gloss.paragraph_format.keep_together = True

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(title)
    format_run(run, persian=True, size=9, bold=True, color=MUTED)
    add_bottom_border(header)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    add_page_field(footer)


def add_cover(doc: Document, reader: dict, teacher: bool) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(62)

    kicker = doc.add_paragraph()
    set_paragraph_rtl(kicker)
    kicker.paragraph_format.space_after = Pt(14)
    run = kicker.add_run("مجموعهٔ خوانش فارسی")
    format_run(run, persian=True, size=12, bold=True, color=ORANGE)

    title = doc.add_paragraph()
    set_paragraph_rtl(title)
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run(reader["title_fa"])
    format_run(run, persian=True, size=30, bold=True, color=GREEN)

    edition = doc.add_paragraph()
    set_paragraph_rtl(edition)
    edition.paragraph_format.space_after = Pt(16)
    run = edition.add_run("نسخهٔ مدرس؛ فارسی و انگلیسی" if teacher else "نسخهٔ فارسی دانشجو")
    format_run(run, persian=True, size=16, bold=True, color=INK)

    if teacher:
        subtitle = doc.add_paragraph()
        set_paragraph_rtl(subtitle, False)
        subtitle.paragraph_format.space_after = Pt(28)
        run = subtitle.add_run(f"{reader['title_en']} — Complete Teacher Reader")
        format_run(run, persian=False, size=13, color=MUTED)

    description = doc.add_paragraph()
    set_paragraph_rtl(description)
    description.paragraph_format.space_before = Pt(18)
    description.paragraph_format.space_after = Pt(18)
    description.paragraph_format.line_spacing = 1.55
    shade_paragraph(description, SOFT)
    text = (
        "این کتاب همهٔ بخش‌های دوره را در یک فایل گرد آورده است. متن فارسی را نخست بخوانید و سپس برای بررسی معنا، از ترجمهٔ انگلیسی و یادداشت‌های مدرس کمک بگیرید."
        if teacher
        else "این کتاب همهٔ بخش‌های دوره را در یک فایل گرد آورده است. متن‌ها را با صدای بلند بخوانید، واژه‌های تازه را نشانه‌گذاری کنید و پس از هر بخش دربارهٔ پرسش پایانی گفت‌وگو کنید."
    )
    run = description.add_run(text)
    format_run(run, persian=True, size=13.5, color=GREEN)

    count = doc.add_paragraph()
    set_paragraph_rtl(count)
    count.paragraph_format.space_before = Pt(24)
    run = count.add_run(f"۲۴ بخش کامل • ۱۴۴ جملهٔ آموزشی")
    format_run(run, persian=True, size=11, bold=True, color=MUTED)

    doc.add_page_break()


def clean_vocab(value: str) -> tuple[str, str]:
    if " — " in value:
        return tuple(value.split(" — ", 1))
    return value, ""


def add_section_heading(doc: Document, week: dict, *, teacher: bool) -> None:
    heading = doc.add_paragraph(style="Heading 1")
    set_paragraph_rtl(heading)
    heading.paragraph_format.page_break_before = False
    run = heading.add_run(f"بخش {week['number']} — {week['section_fa']}")
    format_run(run, persian=True, size=16, bold=True, color=GREEN)
    add_bottom_border(heading, color=SOFT)

    if teacher:
        english_title = next((line.get("gloss", "") for line in week["lines"] if line.get("gloss")), "")
        # The first gloss is a sentence, so use the stored section label from the teaching note only as context.
        if english_title:
            label = doc.add_paragraph(style="English Gloss")
            label.paragraph_format.space_after = Pt(9)
            run = label.add_run(f"Section {week['number']} of 24")
            format_run(run, persian=False, size=9.5, bold=True, color=ORANGE)


def add_student_week(doc: Document, week: dict, index: int) -> None:
    add_section_heading(doc, week, teacher=False)
    for line in week["lines"]:
        paragraph = doc.add_paragraph()
        set_paragraph_rtl(paragraph)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.35
        paragraph.paragraph_format.keep_together = True
        number = paragraph.add_run(f"{line['id']}  ")
        format_run(number, persian=True, size=10, bold=True, color=ORANGE)
        text = paragraph.add_run(line["text"])
        format_run(text, persian=True, size=13.5, color=INK)

    vocab = [clean_vocab(item)[0] for item in week.get("vocab", [])]
    if vocab:
        paragraph = doc.add_paragraph()
        set_paragraph_rtl(paragraph)
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(7)
        label = paragraph.add_run("واژگان کلیدی: ")
        format_run(label, persian=True, size=10.5, bold=True, color=GREEN)
        words = paragraph.add_run("، ".join(vocab))
        format_run(words, persian=True, size=10.5, color=MUTED)

    if index % 2 == 0 and index < 24:
        doc.add_page_break()
    else:
        divider = doc.add_paragraph()
        divider.paragraph_format.space_after = Pt(2)
        add_bottom_border(divider, color=LINE)


def add_teacher_week(doc: Document, week: dict, index: int) -> None:
    add_section_heading(doc, week, teacher=True)
    for line in week["lines"]:
        paragraph = doc.add_paragraph()
        set_paragraph_rtl(paragraph)
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = 1.18
        paragraph.paragraph_format.keep_together = True
        number = paragraph.add_run(f"{line['id']}  ")
        format_run(number, persian=True, size=9.5, bold=True, color=ORANGE)
        text = paragraph.add_run(line["text"])
        format_run(text, persian=True, size=12.25, color=INK)

        gloss = doc.add_paragraph(style="English Gloss")
        gloss.paragraph_format.keep_with_next = False
        run = gloss.add_run(line.get("gloss", ""))
        format_run(run, persian=False, size=9.5, color=MUTED)

    vocab_pairs = [clean_vocab(item) for item in week.get("vocab", [])]
    if vocab_pairs:
        paragraph = doc.add_paragraph()
        set_paragraph_rtl(paragraph, False)
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(4)
        label = paragraph.add_run("Vocabulary: ")
        format_run(label, persian=False, size=9, bold=True, color=GREEN)
        pairs = paragraph.add_run("; ".join(f"{fa} — {en}" if en else fa for fa, en in vocab_pairs))
        format_run(pairs, persian=False, size=9, color=MUTED)

    if week.get("note"):
        note = doc.add_paragraph()
        set_paragraph_rtl(note, False)
        note.paragraph_format.space_before = Pt(2)
        note.paragraph_format.space_after = Pt(4)
        shade_paragraph(note, SOFT)
        label = note.add_run("Teaching note: ")
        format_run(label, persian=False, size=8.5, bold=True, color=GREEN)
        text = note.add_run(week["note"])
        format_run(text, persian=False, size=8.5, color=MUTED)

    if index < 24:
        doc.add_page_break()


def build_document(reader: dict, path: Path, teacher: bool) -> None:
    doc = Document()
    configure_document(doc, reader["title_fa"], teacher)
    add_cover(doc, reader, teacher)
    for index, week in enumerate(reader["weeks"], start=1):
        if teacher:
            add_teacher_week(doc, week, index)
        else:
            add_student_week(doc, week, index)
    doc.save(path)


def main() -> None:
    payload = json.loads(DATA.read_text(encoding="utf-8"))
    readers_by_slug = {reader["slug"]: reader for reader in payload["readers"]}
    DOWNLOADS.mkdir(parents=True, exist_ok=True)

    for slug, (student_name, teacher_name) in COURSES.items():
        reader = readers_by_slug[slug]
        reader["student"] = student_name
        reader["teacher"] = teacher_name
        build_document(reader, DOWNLOADS / student_name, teacher=False)
        build_document(reader, DOWNLOADS / teacher_name, teacher=True)
        print(f"Created {student_name} and {teacher_name}")

    DATA.write_text(
        json.dumps(payload, ensure_ascii=False, separators=(",", ":")),
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()
