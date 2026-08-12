from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DATA = ROOT / "data" / "readers.json"
DOWNLOADS = ROOT / "public" / "downloads"
STUDENT = DOWNLOADS / "AFH1_2025_Farsi_Student_Study_Handbook.docx"
TEACHER = DOWNLOADS / "AFH1_2025_Farsi_English_Teacher_Handbook.docx"
PERSIAN_ID = re.compile(r"^([۰-۹]{3}٫[۰-۹]{2})\s{2}")
ASCII_ID = re.compile(r"^(\d{3}\.\d{2})\s{2}")
PERSIAN_TO_ASCII = str.maketrans("۰۱۲۳۴۵۶۷۸۹٫", "0123456789.")


def set_text(paragraph, value: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)


def data_lines():
    payload = json.loads(DATA.read_text(encoding="utf-8"))
    reader = next(reader for reader in payload["readers"] if reader["slug"] == "afh1-2025")
    lines = [line for week in reader["weeks"] for line in week["lines"]]
    by_fa_id = {line["id"]: line["text"] for line in lines if line.get("id")}
    headings = [line["text"] for line in lines if not line.get("id")]
    return by_fa_id, headings


def update_student(by_fa_id: dict[str, str], headings: list[str]) -> tuple[int, int]:
    document = Document(STUDENT)
    sentence_count = 0
    heading_count = 0
    for paragraph in document.paragraphs:
        value = paragraph.text.strip()
        if not value or value.startswith("هفتهٔ"):
            continue
        match = PERSIAN_ID.match(value)
        if match:
            identifier = match.group(1)
            set_text(paragraph, f"{identifier}  {by_fa_id[identifier]}")
            sentence_count += 1
        else:
            if heading_count >= len(headings):
                raise RuntimeError("student handbook contains more unnumbered headings than the reader data")
            set_text(paragraph, headings[heading_count])
            heading_count += 1
    if sentence_count != len(by_fa_id) or heading_count != len(headings):
        raise RuntimeError(
            f"student handbook alignment failed: {sentence_count}/{len(by_fa_id)} sentences, "
            f"{heading_count}/{len(headings)} headings"
        )
    document.save(STUDENT)
    return sentence_count, heading_count


def update_teacher(by_fa_id: dict[str, str]) -> int:
    document = Document(TEACHER)
    sentence_count = 0
    by_ascii_id = {identifier.translate(PERSIAN_TO_ASCII): text for identifier, text in by_fa_id.items()}
    for paragraph in document.paragraphs:
        match = ASCII_ID.match(paragraph.text.strip())
        if not match:
            continue
        identifier = match.group(1)
        set_text(paragraph, f"{identifier}  {by_ascii_id[identifier]}")
        sentence_count += 1
    if sentence_count != len(by_ascii_id):
        raise RuntimeError(f"teacher handbook alignment failed: {sentence_count}/{len(by_ascii_id)} sentences")
    document.save(TEACHER)
    return sentence_count


def main() -> None:
    by_fa_id, headings = data_lines()
    student_sentences, student_headings = update_student(by_fa_id, headings)
    teacher_sentences = update_teacher(by_fa_id)
    print(json.dumps({
        "student_sentences": student_sentences,
        "student_headings": student_headings,
        "teacher_sentences": teacher_sentences,
    }))


if __name__ == "__main__":
    main()
